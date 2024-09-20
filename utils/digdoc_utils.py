import os
import re
import datetime

import ebooklib
from ebooklib import epub
import mobi
import docx
import nbformat

from bs4 import BeautifulSoup

from langchain.text_splitter import CharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters.base import Document

from langchain_openai import OpenAIEmbeddings
from langchain_openai import ChatOpenAI
from langchain_anthropic import ChatAnthropic

from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate

from .utils import get_context, save_response
from .WebScraper import WebScraper

import numpy as np
from rank_bm25 import BM25Okapi
from sklearn.metrics.pairwise import cosine_similarity

from dotenv import load_dotenv
load_dotenv()

BASE_DIREC = "projects"

class DigDoc:

    def __init__(self, **params):
        self.model = params.get("model", "gpt-4o")
        self.vectorstore = None

        self.target = params.get("target", [])
        self.reindex = params.get("reindex", True)

        self.project_name = params.get("project_name", "default")
        self.directory = params.get("directory", "web")
        self.file_name = os.path.join(BASE_DIREC, self.project_name, "ChatHistory.html")
        os.makedirs(BASE_DIREC, exist_ok=True)
        self.project_path = os.path.join(BASE_DIREC, self.project_name)
        os.makedirs(self.project_path, exist_ok=True)

        self.look_back_window = 3
        self.chat_history = []

        self.max_depth = 1
        self.max_pages = 1
        self.load_from_cache = False

    def answer(self, query):

        if self.directory == "google":
            self.dig(query)

        # Create a custom prompt template
        prompt_template = """
        Your role is a research assistant at a university that helps with reviewing documents.
        You have been asked to find information based on the provided context and question.
        Say you don't know if there are not enough related information\n"""

        prompt_template = "[HISTORY STARTS] Here is the chat history:\n"
        prompt_template += f"{get_context(self.chat_history, self.look_back_window)}\n"
        prompt_template += "[HISTORY ENDS]"

        prompt_template += """
        Context:
        {context}

        Question: {question}
        Answer: """

        PROMPT = PromptTemplate(
            template=prompt_template,
            input_variables=["context", "question"]
        )
        
        # Retrieve Top-N candidates using BM25
        query_terms = query.split()
        bm25_indices = [self.original_texts.index(doc) for doc in self.bm25.get_top_n(query_terms, self.original_texts, n=100)]
        bm25_candidates = [self.original_texts[idx] for idx in bm25_indices]

        # Use your embeddings model to compute the embeddings
        embeddings_model = OpenAIEmbeddings()  # Change this to your specific embeddings model if different

        # Compute embeddings for the query
        query_embedding = embeddings_model.embed_query(query)
        query_embedding = np.array(query_embedding).reshape(1, -1)

        # Compute embeddings for BM25 candidates
        candidate_embeddings = embeddings_model.embed_documents(bm25_candidates)
        candidate_embeddings = np.array(candidate_embeddings)


        # Calculate cosine similarity
        similarities = cosine_similarity(query_embedding, candidate_embeddings)[0]

        # Select top K documents based on cosine similarity
        top_k_indices = np.argsort(similarities)[-5:][::-1]
        top_k_docs = [bm25_candidates[i] for i in top_k_indices]

        # Compose final context based on top K documents
        context = ' '.join(top_k_docs)

        llm = get_llm(self.model)
        response = llm.invoke(f"{prompt_template.format(context=context, question=query)}").content


        # # Ensure necessary keys are present in the result
        # if 'result' not in result or 'source_documents' not in result:
        #     raise ValueError("The response from qa_chain is missing expected keys.")

        # response, source_docs = result["result"], result["source_documents"]

        self.chat_history.append({
            "query": query,
            "response": response
        })

        save_response(self.chat_history, self.file_name, self.project_name)

        # return result["result"], result["source_documents"]
        return response

    def dig(self, query=None):
        if not query and self.directory=="google":
            """
            The reason is that we need to have a way to pass the query to the dig method
            The Jupyter architecture cannot be changed, so I will dig when the answer method is called with the query
            """
            return

        if self.directory == "google":
            assert len(self.target) == 1, "You need to pass a single website, or 'all' to the target if directory is google"
            assert self.max_depth == 1, "Max depth should be 1 for google search"
            assert self.max_pages <= 10, "Max pages should be max 10 for google search"

        if len(self.target) == 0:
            raise ValueError("You need to pass at least one document to read.")

        # This is the case when we have already vectorized the documents
        if os.path.exists(self.project_path) and not self.reindex:
            x = FAISS.load_local(
                self.project_path,
                embeddings=OpenAIEmbeddings(),
                allow_dangerous_deserialization=True
            )
            self.vectorstore = x

        # Iterate over all PDFs in the directory recursively in all folders
        if self.directory == "web":
            raw_documents = self.get_content_of_website()

        elif self.directory == "google":
            raw_documents = self.get_content_from_google_search(query)

        else:
            raw_documents = self.get_documents_raw_text()

        # Split text into chunks
        text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=0)
        texts = text_splitter.split_documents(raw_documents)

        # Create embeddings
        # embeddings = OpenAIEmbeddings()

        # Create vector store
        # vectorstore = FAISS.from_documents(texts, embeddings)
        # vectorstore.save_local(self.project_path)

        # self.vectorstore = vectorstore

        # Store texts separately for BM25
        self.original_texts = [doc.page_content for doc in texts]


        # Initialize BM25 retriever
        tokenized_texts = [doc.page_content.split() for doc in texts]  # Tokenize the documents for BM25
        self.bm25 = BM25Okapi(tokenized_texts)

    def set_scrapping_params(self, max_depth, max_pages, load_from_cache):
        self.max_depth = max_depth
        self.max_pages = max_pages
        self.load_from_cache = load_from_cache

    def get_documents_raw_text(self):
        raw_documents = []
        n_files = 0

        for root, dirs, files in os.walk(self.directory):
            for filename in files:

                if isinstance(self.target, list) and \
                        not any(doc.lower() in filename.lower() for doc in self.target):
                    continue

                print (filename.lower(), "is being processed")

                if filename.lower().endswith('.pdf'):
                    pdf_path = os.path.join(root, filename)
                    loader = PyPDFLoader(pdf_path)
                    raw_documents.extend(loader.load())
                    n_files += 1

                elif filename.lower().endswith('.txt'):
                    with open(os.path.join(root, filename), 'r') as f:
                        file_content = f.read()
                    raw_documents.append(Document(file_content))
                    n_files += 1

                elif filename.lower().endswith('.epub'):
                    book = epub.read_epub(os.path.join(root, filename))
                    for item in book.get_items():
                        if item.get_type() == ebooklib.ITEM_DOCUMENT:
                            content = item.get_content()
                            if isinstance(content, bytes):
                                content = content.decode('utf-8')

                            text = BeautifulSoup(content, 'html.parser').get_text(strip=True)
                            if len(text) < 100:
                                continue

                            raw_documents.append(Document(content))
                            n_files += 1

                # Add cpp and c files
                elif filename.lower().split('.')[-1] in ['cpp', 'c', 'h']:
                    with open(os.path.join(root, filename), 'r') as f:
                        file_content = f.read()
                    raw_documents.append(Document(file_content))
                    n_files += 1

                elif filename.lower().endswith(".mobi"):
                    book = mobi.Mobi(os.path.join(root, filename))
                    content = book.get_text()
                    if content:
                        raw_documents.append(Document(content))
                        n_files += 1

                elif filename.lower().endswith(".docx") or filename.endswith(".doc"):
                    doc = docx.Document(os.path.join(root, filename))
                    content = ""
                    for para in doc.paragraphs:
                        content += para.text + "\n"
                    raw_documents.append(Document(content))
                    n_files += 1

                elif filename.lower().endswith(".py"):
                    with open(os.path.join(root, filename), 'r') as f:
                        file_content = f.read()
                    raw_documents.append(Document(file_content))
                    n_files += 1

                elif filename.lower().endswith(".ipynb"):
                    with open(os.path.join(root, filename), 'r') as f:
                        file_content = f.read()
                    nb = nbformat.reads(file_content, as_version=4)
                    for cell in nb.cells:
                        if cell.cell_type == 'code':
                            raw_documents.append(Document(cell.source))
                    n_files += 1

                else:
                    pass

        if n_files == 0:
            raise ValueError("No Files found.")

        return raw_documents
    
    def get_content_of_website(self):

        raw_documents = []
        scraper = WebScraper(
            max_pages=self.max_pages,
            max_depth=self.max_depth,
            load_from_cache=self.load_from_cache,
            project_path=self.project_path
        )

        for url in self.target:

            contents = scraper.scrape_website(url)
            for content in contents:

                new_doc = Document(content[1])
                new_doc.metadata['source'] = content[0]
                raw_documents.append(new_doc)

        return raw_documents

    def get_content_from_google_search(self, query):

        raw_documents = []
        scraper = WebScraper(
            max_pages=self.max_pages,
            max_depth=self.max_depth,
            load_from_cache=self.load_from_cache,
            project_path=self.project_path
        )

        contents = scraper.scrape_google_search(self.target[0], query)
        for url, content in contents:

            new_doc = Document(content)
            new_doc.metadata['source'] = url
            raw_documents.append(new_doc)

        return raw_documents


# ---------------------------------------------------------------------------- #
#                       Utils for Digestor                                    #
# ---------------------------------------------------------------------------- #

def generate_markdown_output(question, answer, source_docs):
    # Remove the HISTORY STARTS and HISTORY ENDS and anything in between from
    # the question and answer
    question = remove_between_tags(question, "[HISTORY STARTS]", "[HISTORY ENDS]")

    markdown = f"""

# {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}

## Answer
--------------------------------------------------------------------------------
{answer}
--------------------------------------------------------------------------------

## Question
{question}

## Source Documents
"""
    for i, doc in enumerate(source_docs):
        source = doc.metadata.get('source', 'Unknown')
        page = doc.metadata.get('page', 'Unknown')
        content = doc.page_content[:200] + "..." if len(doc.page_content) > 200 else doc.page_content

        markdown += f"""
### Document {i + 1}
- **Source**: {source}
- **Page**: {page}

#### Content Preview {content}

---
"""
    return markdown


def save_markdown(markdown, project_name, append=False):
    file_path = os.path.join("../projects", project_name, "ChatHistory.md")

    # Ensure the file exists, create it if it doesn't
    if not os.path.exists(file_path):
        open(file_path, 'w', encoding="utf-8").close()

    existing_content = ""

    # Read existing content
    if append:
        try:
            with open(file_path, 'r', encoding="utf-8") as file:
                existing_content = file.read()
        except UnicodeDecodeError:
            # If utf-8 fails, try with a different encoding
            with open(file_path, 'r', encoding="iso-8859-1") as file:
                existing_content = file.read()

    # Write new content
    with open(file_path, 'w', encoding="utf-8") as file:
        file.write(markdown + existing_content)

    print(f"Markdown saved to {file_path}")


def remove_between_tags(text, start_tag, end_tag):
    pattern = f'{re.escape(start_tag)}.*?{re.escape(end_tag)}'
    return re.sub(pattern, '', text, flags=re.DOTALL)


def get_llm(model):
    if model == "gpt-4o":
        return ChatOpenAI(model_name="gpt-4o")

    elif model == "anthropic":
        return ChatAnthropic(model_name="claude-3-5-sonnet-20240620")

    else:
        raise ValueError(f"Model {model} not found.")


