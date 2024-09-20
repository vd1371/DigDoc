import os

import ebooklib
from ebooklib import epub
import mobi
import docx
import nbformat
from bs4 import BeautifulSoup

from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters.base import Document

from .WebScraper import WebScraper

class DocumentReader:

    def __init__(self, **params):

        self.directory = params['directory']
        self.target = params['target']
        self.max_depth = params['max_depth']
        self.max_pages = params['max_pages']
        self.load_from_cache = params['load_from_cache']
        self.project_path = params['project_path']


        if self.directory == "google":
            assert len(self.target) == 1, "You need to pass a single website, or 'all' to the target if directory is google"
            assert self.max_depth == 1, "Max depth should be 1 for google search"
            assert self.max_pages <= 10, "Max pages should be max 10 for google search"

        if len(self.target) == 0:
            raise ValueError("You need to pass at least one document to read.")


    def read(self, query=None):

        # Iterate over all PDFs in the directory recursively in all folders
        if self.directory == "web":
            raw_documents = self.get_content_of_website()

        elif self.directory == "google":
            raw_documents = self.get_content_from_google_search(query)

        else:
            raw_documents = self.get_documents_raw_text()

        return raw_documents

    

    def get_documents_raw_text(self):
        raw_documents = []
        n_files = 0

        for root, dirs, files in os.walk(self.directory):
            for filename in files:

                if isinstance(self.target, list) and \
                        not any(doc.lower() in filename.lower() for doc in self.target):
                    continue

                print (filename.lower(), "is being processed")

                if filename.lower().endswith('.pdf'):
                    pdf_path = os.path.join(root, filename)
                    loader = PyPDFLoader(pdf_path)
                    raw_documents.extend(loader.load())
                    n_files += 1

                elif filename.lower().endswith('.txt'):
                    with open(os.path.join(root, filename), 'r') as f:
                        file_content = f.read()
                    raw_documents.append(Document(file_content))
                    n_files += 1

                elif filename.lower().endswith('.epub'):
                    book = epub.read_epub(os.path.join(root, filename))
                    for item in book.get_items():
                        if item.get_type() == ebooklib.ITEM_DOCUMENT:
                            content = item.get_content()
                            if isinstance(content, bytes):
                                content = content.decode('utf-8')

                            text = BeautifulSoup(content, 'html.parser').get_text(strip=True)
                            if len(text) < 100:
                                continue

                            raw_documents.append(Document(content))
                            n_files += 1

                # Add cpp and c files
                elif filename.lower().split('.')[-1] in ['cpp', 'c', 'h']:
                    with open(os.path.join(root, filename), 'r') as f:
                        file_content = f.read()
                    raw_documents.append(Document(file_content))
                    n_files += 1

                elif filename.lower().endswith(".mobi"):
                    book = mobi.Mobi(os.path.join(root, filename))
                    content = book.get_text()
                    if content:
                        raw_documents.append(Document(content))
                        n_files += 1

                elif filename.lower().endswith(".docx") or filename.endswith(".doc"):
                    doc = docx.Document(os.path.join(root, filename))
                    content = ""
                    for para in doc.paragraphs:
                        content += para.text + "\n"
                    raw_documents.append(Document(content))
                    n_files += 1

                elif filename.lower().endswith(".py"):
                    with open(os.path.join(root, filename), 'r') as f:
                        file_content = f.read()
                    raw_documents.append(Document(file_content))
                    n_files += 1

                elif filename.lower().endswith(".ipynb"):
                    with open(os.path.join(root, filename), 'r') as f:
                        file_content = f.read()
                    nb = nbformat.reads(file_content, as_version=4)
                    for cell in nb.cells:
                        if cell.cell_type == 'code':
                            raw_documents.append(Document(cell.source))
                    n_files += 1

                else:
                    pass

        if n_files == 0:
            raise ValueError("No Files found.")

        return raw_documents
    
    def get_content_of_website(self):

        raw_documents = []
        scraper = WebScraper(
            max_pages=self.max_pages,
            max_depth=self.max_depth,
            load_from_cache=self.load_from_cache,
            project_path=self.project_path
        )

        for url in self.target:

            contents = scraper.scrape_website(url)
            for content in contents:

                new_doc = Document(content[1])
                new_doc.metadata['source'] = content[0]
                raw_documents.append(new_doc)

        return raw_documents

    def get_content_from_google_search(self, query):

        raw_documents = []
        scraper = WebScraper(
            max_pages=self.max_pages,
            max_depth=self.max_depth,
            load_from_cache=self.load_from_cache,
            project_path=self.project_path
        )

        contents = scraper.scrape_google_search(self.target[0], query)
        for url, content in contents:

            new_doc = Document(content)
            new_doc.metadata['source'] = url
            raw_documents.append(new_doc)

        return raw_documents

