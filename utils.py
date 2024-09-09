import os
import re

from openai import OpenAI
import anthropic
from anthropic.types import TextBlock

import ebooklib
from ebooklib import epub
import mobi
import docx

from langchain.text_splitter import CharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.document_loaders import PyPDFLoader
from langchain_text_splitters.base import Document

from langchain_openai import ChatOpenAI
from langchain_openai import OpenAIEmbeddings

from langchain_anthropic import ChatAnthropic

from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate

from dotenv import load_dotenv

load_dotenv()


class Digestor:

    def __init__(self, docs_directory=None, model="gpt-4o", project_name="literature_review"):
        self.model = model
        self.vectorstore = None
        self.chat_history = []
        self.project_name = project_name
        self.docs_directory = docs_directory

    def answer(self, question, refresh_history=False):
        # Create a custom prompt template
        prompt_template = """
        Your role is a research assistant at a university that helps with reviewing documents.
        You have been asked to find information based on the provided context and question.
        Use the following pieces of context to answer the question.\n"""

        prompt_template += f"{convert_chat_history(self.chat_history, n_turns=2, refresh=refresh_history)}"

        prompt_template += """
        Context:
        {context}

        Question: {question}
        Answer: """

        PROMPT = PromptTemplate(
            template=prompt_template,
            input_variables=["context", "question"]
        )

        # Create a retrieval-based question-answering chain with GPT-4
        qa_chain = RetrievalQA.from_chain_type(
            llm=get_llm(self.model),
            chain_type="stuff",
            retriever=self.vectorstore.as_retriever(search_kwargs={"k": 5}),  # Retrieve top 4 documents
            return_source_documents=True,
            chain_type_kwargs={"prompt": PROMPT}
        )

        result = qa_chain.invoke({
            "query": question,
        })

        # Ensure necessary keys are present in the result
        if 'result' not in result or 'source_documents' not in result:
            raise ValueError("The response from qa_chain is missing expected keys.")

        answer, source_docs = result["result"], result["source_documents"]
        markdown = generate_markdown_output(question, answer, source_docs)
        save_markdown(markdown, self.project_name, append=True)

        self.chat_history.append({
            "question": question,
            "answer": answer
        })

        return result["result"], result["source_documents"]

    def read_docs(self, target_docs=None, reindex=False):
        """
        :param target_docs: list of strings to filter the PDFs to vectorize
        :param reindex: whether to reindex the documents
        """

        if isinstance(target_docs, list) and len(target_docs) == 0:
            target_docs = "all"

        base_direc = "projects"
        os.makedirs(base_direc, exist_ok=True)
        project_path = os.path.join(base_direc, self.project_name)

        # This is the case when we have already vectorized the documents
        if os.path.exists(project_path) and not reindex and target_docs == "all":
            x = FAISS.load_local(
                project_path,
                embeddings=OpenAIEmbeddings(),
                allow_dangerous_deserialization=True
            )
            self.vectorstore = x

        # Iterate over all PDFs in the directory recursively in all folders
        raw_documents = get_documents_raw_text(self.docs_directory, target_docs)

        # Split text into chunks
        text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=0)
        texts = text_splitter.split_documents(raw_documents)

        # Create embeddings
        embeddings = OpenAIEmbeddings()

        # Create vector store
        vectorstore = FAISS.from_documents(texts, embeddings)
        vectorstore.save_local(project_path)

        self.vectorstore = vectorstore


# ---------------------------------------------------------------------------- #
#                       Utils for Digestor                                    #
# ---------------------------------------------------------------------------- #

def convert_chat_history(chat_history, n_turns=2, refresh=False):
    """
    :param chat_history: list of dicts with keys 'question', 'answer'
    :return: string with chat history formatted as a conversation
    """

    if refresh:
        return ""

    if len(chat_history) == 0:
        return ""

    chat_history = chat_history[-n_turns:]
    chat = "[HISTORY STARTS] Here is the chat history:\n"
    for i, chat_his in enumerate(chat_history):
        chat += f"User: {chat_his['question']}\n"
        chat += f"AI: {chat_his['answer']}\n"

    chat += "[HISTORY ENDS]"

    return chat


def generate_markdown_output(question, answer, source_docs):
    # Remove the HISTORY STARTS and HISTORY ENDS and anything in between from
    # the question and answer
    question = remove_between_tags(question, "[HISTORY STARTS]", "[HISTORY ENDS]")

    markdown = f"""

## Answer
--------------------------------------------------------------------------------
{answer}
--------------------------------------------------------------------------------

## Question
{question}

## Source Documents
"""
    for i, doc in enumerate(source_docs):
        source = doc.metadata.get('source', 'Unknown')
        page = doc.metadata.get('page', 'Unknown')
        content = doc.page_content[:200] + "..." if len(doc.page_content) > 200 else doc.page_content

        markdown += f"""
### Document {i + 1}
- **Source**: {source}
- **Page**: {page}

#### Content Preview {content}

---
"""
    return markdown


def save_markdown(markdown, project_name, append=False):
    file_path = os.path.join("projects", project_name, "ChatHistory.md")

    # Ensure the file exists, create it if it doesn't
    if not os.path.exists(file_path):
        open(file_path, 'w', encoding="utf-8").close()

    existing_content = ""

    # Read existing content
    if append:
        try:
            with open(file_path, 'r', encoding="utf-8") as file:
                existing_content = file.read()
        except UnicodeDecodeError:
            # If utf-8 fails, try with a different encoding
            with open(file_path, 'r', encoding="iso-8859-1") as file:
                existing_content = file.read()

    # Write new content
    with open(file_path, 'w', encoding="utf-8") as file:
        file.write(markdown + existing_content)

    print(f"Markdown saved to {file_path}")


def remove_between_tags(text, start_tag, end_tag):
    pattern = f'{re.escape(start_tag)}.*?{re.escape(end_tag)}'
    return re.sub(pattern, '', text, flags=re.DOTALL)


def get_llm(model):
    if model == "gpt-4o":
        return ChatOpenAI(model_name="gpt-4o")

    elif model == "anthropic":
        return ChatAnthropic(model_name="claude-3-5-sonnet-20240620")

    else:
        raise ValueError(f"Model {model} not found.")


def get_documents_raw_text(directory, target_docs):
    raw_documents = []
    n_files = 0

    for root, dirs, files in os.walk(directory):
        for filename in files:

            if isinstance(target_docs, list) and \
                    not any(doc.lower() in filename.lower() for doc in target_docs):
                continue

            if filename.endswith('.pdf'):

                print(filename, "is being processed", end="\r")

                pdf_path = os.path.join(root, filename)
                loader = PyPDFLoader(pdf_path)
                raw_documents.extend(loader.load())

                n_files += 1

            elif filename.endswith('.txt'):
                with open(os.path.join(root, filename), 'r') as f:
                    file_content = f.read()
                raw_documents.append(Document(file_content))

                n_files += 1

            elif filename.endswith('.epub'):
                book = epub.read_epub(os.path.join(root, filename))
                for item in book.get_items():
                    if item.get_type() == ebooklib.ITEM_DOCUMENT:

                        content = item.get_content()
                        # Convert to string
                        if isinstance(content, bytes):
                            content = content.decode('utf-8')
                        raw_documents.append(Document(content))

                        n_files += 1

            elif filename.endswith(".mobi"):
                book = mobi.Mobi(os.path.join(root, filename))

                content = book.get_text()
                if content:
                    raw_documents.append(Document(content))
                    n_files += 1

            elif filename.endswith(".docx"):
                doc = docx.Document(os.path.join(root, filename))
                content = ""
                for para in doc.paragraphs:
                    content += para.text + "\n"
                raw_documents.append(Document(content))
                n_files += 1

            else:
                pass

    if n_files == 0:
        raise ValueError("No Files found.")

    return raw_documents


# ---------------------------------------------------------------------------- #
#                       Utils for HER model                                    #
# ---------------------------------------------------------------------------- #

class Her:

    def __init__(self, model="gpt-4o", should_refresh_context=False, append_history=False):
        self.model = model
        self.should_refresh_context = should_refresh_context
        self.append_history = append_history

    def ask(self, query, role=None):
        if role is None:
            role = "You are a general assistant helping with various tasks."

        if self.model == "gpt-4o":
            client = OpenAI(
                api_key=os.getenv("API_KEY"),
            )
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": role},
                    {"role": "user",
                     "content": f"Context: {get_context_for_her(self.should_refresh_context)}\n\nQuestion: {query}"}
                ]
            )
            response = response.choices[0].message.content

        elif self.model == "anthropic":
            client = anthropic.Anthropic(
                api_key=os.getenv("ANTHROPIC_API_KEY"),
            )
            message = client.messages.create(
                model="claude-3-5-sonnet-20240620",
                max_tokens=4096,
                messages=[
                    {"role": "user",
                     "content": f"Context: {get_context_for_her(self.should_refresh_context)}\n\nQuestion: {query}"}
                ]
            )

            # Extract text from TextBlock
            if isinstance(message.content[0], TextBlock):
                response = message.content[0].text
            else:
                response = str(message.content[0])

        save_response_to_her(response, append=self.append_history)

        return response


def get_context_for_her(refresh=False):
    if refresh:
        return ""

    # Open a markdown file and write the response
    file_name = "Her.md"
    if not os.path.exists(file_name):
        with open(file_name, "w") as f:
            f.write("")

    with open(file_name, "r") as f:
        context = f.read()

    return context


def save_response_to_her(response, file_name="TalkWithHer.md", append=False):
    to_write = response
    try:
        with open(file_name, "w") as f:
            current_file = f.read()
    except Exception as e:
        current_file = ""

    if append:
        to_write += "\n\n" + current_file

    with open(file_name, "w") as f:
        f.write(to_write)

    print(f"Response saved to {file_name}")
